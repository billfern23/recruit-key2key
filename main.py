import nltk.corpus
import numpy as np
import docx2txt
import re
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import utility.keyword as ut

punctuation = re.compile(r'[-.?!@,:;()|0-9]')
# creating files
resume = docx2txt.process("Resume2.docx")
cv = docx2txt.process("CV.docx")
job_description = docx2txt.process("Job_Description.docx")

# A list of text
text = [resume, job_description]
# Create a count vectorizer object and convert the text documents to a matrix of token counts
count_vec = CountVectorizer()
count_matrix = count_vec.fit_transform(text)
# Get and print the cosine similarity scores
print("\nSimilarity Scores:")
print(cosine_similarity(count_matrix))

# Get the match percentage
matchPercentage = cosine_similarity(count_matrix)[0][1] * 100
matchPercentage = round(matchPercentage, 2) # round to two decimal
print("Your resume matches about "+ str(matchPercentage)+ "% of the job description.")

data = ut.find_keyword("flexible, creative, analytical, organized, quick learner", resume)
print(data)

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
doc = Document('resume.docx')
lists = ["Range", "adaptible", "quick learner", "essential"]  

for paragraph in doc.paragraphs:
    for target in lists:
        if target in paragraph.text:  

            currRuns = copy.copy(paragraph.runs)   
            paragraph.runs.clear()

            for run in currRuns:
                if target in run.text:
                    words = re.split('(\W)', run.text)  
                    for word in words:
                        if word == target:
                            newRun = paragraph.add_run(word)
                            newRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            newRun = paragraph.add_run(word)
                            newRun.font.highlight_color = None
                else: # 
                    paragraph.runs.append(run)


doc.save('resume1.docx')
