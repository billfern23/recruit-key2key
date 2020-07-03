from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
doc = Document('resume.docx')
keywords = ["Flexible", "quick", "genius"]  

for paragraph in doc.paragraphs:
    for target in keywords:
        if target in paragraph.text:  

            Checker = copy.copy(paragraph.runs)   
            paragraph.clear()

            for run in Checker:
                if target in run.text:
                    words = re.split('(\W)', run.text)  
                    for word in words:
                        if word == target:
                            highlights = paragraph.add_run(word)
                            highlights.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            highlights = paragraph.add_run(word)
                            highlights.font.highlight_color = None
                else: 
                    paragraph.runs.append(run)


doc.save('resume1.docx')
